from flask import Flask, jsonify, request
import requests
import os
import hashlib
import json
import time
from blockchain import Blockchain
from web3 import Web3
from eth_account import Account
from p2p import NodeRegistry
import sys
import socket
import pdfplumber
from docx import Document
import io
from datasketch import MinHash
from eth_abi import decode
import re
import unicodedata
from flask_cors import CORS

def get_local_ip():
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:
        s.connect(('8.8.8.8', 1))
        ip = s.getsockname()[0]
    except:
        ip = '127.0.0.1'
    finally:
        s.close()
    return ip

bootstrap_url = "http://192.168.1.8:5000"
node_registry = NodeRegistry(bootstrap_url=bootstrap_url)

# Kết nối với Ganache
web3 = Web3(Web3.HTTPProvider("http://127.0.0.1:8545"))

# Địa chỉ contract
contract_address = "0x5FbDB2315678afecb367f032d93F642f64180aa3"

# ABI của smart contract
contract_abi = [
    {
        "anonymous": False,
        "inputs": [{"indexed": False, "internalType": "string", "name": "documentHash", "type": "string"}],
        "name": "DocumentStored",
        "type": "event"
    },
    {
        "inputs": [{"internalType": "string", "name": "documentHash", "type": "string"}],
        "name": "storeDocument",
        "outputs": [],
        "stateMutability": "nonpayable",
        "type": "function"
    },
    {
        "inputs": [{"internalType": "string", "name": "documentHash", "type": "string"}],
        "name": "verifyDocument",
        "outputs": [{"internalType": "bool", "name": "", "type": "bool"}],
        "stateMutability": "view",
        "type": "function"
    }
]

# Load contract
contract = web3.eth.contract(address=contract_address, abi=contract_abi)

if web3.is_connected():
    print("✅ Đã kết nối với Ethereum node!")
else:
    print("❌ Không thể kết nối với Ethereum node!")

app = Flask(__name__)
CORS(app)
blockchain = Blockchain()

# Đảm bảo thư mục lưu trữ file
UPLOAD_FOLDER = 'Uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Danh sách từ khóa không lành mạnh
BLACKLIST_KEYWORDS = ['offensive', 'inappropriate', 'hate', 'violence', 'illegal']

def broadcast_with_retry(url, data=None, files=None, retries=3, timeout=60):
    for attempt in range(retries):
        try:
            if files:
                response = requests.post(url, data=data, files=files, timeout=timeout)
            else:
                response = requests.post(url, json=data, timeout=timeout)
            print(f"Broadcast tới {url}: {response.status_code}, response={response.text}")
            return response
        except Exception as e:
            print(f"Lỗi khi broadcast tới {url}, lần thử {attempt + 1}/{retries}: {str(e)}")
            if attempt < retries - 1:
                time.sleep(2)
    print(f"Broadcast tới {url} thất bại sau {retries} lần thử")
    return None

def check_node_status(node_url, retries=5, delay=2):
    for attempt in range(retries):
        try:
            response = requests.get(f'{node_url}/ping', timeout=5)
            if response.status_code == 200:
                print(f"Node {node_url} sẵn sàng")
                return True
        except Exception as e:
            print(f"Node {node_url} chưa sẵn sàng, lần thử {attempt + 1}/{retries}: {str(e)}")
            time.sleep(delay)
    print(f"Node {node_url} không sẵn sàng sau {retries} lần thử")
    return False

def normalize_text(text):
    """Chuẩn hóa văn bản trước khi tạo MinHash."""
    if not text:
        return ""
    text = unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('ascii')
    text = text.lower()
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\w\s]', '', text)
    return text.strip()

def extract_text(file_content, filename):
    """Trích xuất văn bản từ file."""
    text_extensions = ['.txt', '.md']
    supported_extensions = text_extensions + ['.pdf', '.docx']
    
    if any(filename.lower().endswith(ext) for ext in supported_extensions):
        try:
            text = ""
            if filename.lower().endswith('.txt') or filename.lower().endswith('.md'):
                text = file_content.decode('utf-8', errors='ignore')
            elif filename.lower().endswith('.pdf'):
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        text += page_text
            elif filename.lower().endswith('.docx'):
                doc = Document(io.BytesIO(file_content))
                for para in doc.paragraphs:
                    text += para.text + "\n"
            text = normalize_text(text)
            if not text:
                print(f"Không trích xuất được văn bản từ {filename}: Văn bản rỗng hoặc chỉ chứa khoảng trắng")
                return None
            print(f"Trích xuất văn bản từ {filename}: {text[:50]}...")
            return text
        except Exception as e:
            print(f"Lỗi khi trích xuất văn bản từ {filename}: {str(e)}")
            return None
    print(f"Định dạng file {filename} không được hỗ trợ")
    return None

def check_text_content(file_content, filename):
    """Kiểm tra nội dung văn bản để phát hiện từ khóa không lành mạnh."""
    text = extract_text(file_content, filename)
    if text:
        for keyword in BLACKLIST_KEYWORDS:
            if keyword.lower() in text.lower():
                print(f"Phát hiện từ khóa không lành mạnh trong {filename}: {keyword}")
                return False, f"Nội dung chứa từ khóa không lành mạnh: {keyword}"
        return True, "Nội dung văn bản hợp lệ"
    return False, "File rỗng hoặc chỉ chứa khoảng trắng"

def get_content_hash(file_content, filename):
    """Tạo MinHash từ nội dung văn bản."""
    text = extract_text(file_content, filename)
    if text:
        words = text.split()
        shingles = {f"{words[i]} {words[i+1]} {words[i+2]}" for i in range(len(words)-2)}
        m = MinHash(num_perm=128)
        for s in shingles:
            m.update(s.encode('utf8'))
        hash_value = ','.join(map(str, m.hashvalues))
        print(f"Tạo content_hash cho {filename}: {hash_value[:50]}...")
        return hash_value
    print(f"Không tạo được content_hash cho {filename}")
    return None

# def jaccard_similarity(hash1, hash2):
#     """Tính Jaccard similarity giữa hai MinHash."""
#     if not hash1 or not hash2:
#         raise ValueError(f"Jaccard similarity: Một trong hai hash rỗng - hash1={hash1}, hash2={hash2}")
#     try:
#         h1 = [int(x) for x in hash1.split(',')]
#         h2 = [int(x) for x in hash2.split(',')]
#         if len(h1) != 128 or len(h2) != 128:
#             raise ValueError(f"Hash không hợp lệ: len(hash1)={len(h1)}, len(hash2)={len(h2)}")
#         m1, m2 = MinHash(num_perm=128), MinHash(num_perm=128)
#         m1.hashvalues = h1
#         m2.hashvalues = h2
#         similarity = m1.jaccard(m2)
#         print(f"Jaccard similarity giữa hash1={hash1[:50]}... và hash2={hash2[:50]}...: {similarity:.4f}")
#         return similarity
#     except Exception as e:
#         print(f"Lỗi nghiêm trọng khi tính Jaccard similarity: hash1={hash1[:50]}..., hash2={hash2[:50]}..., error={str(e)}")
#         raise

def jaccard_similarity(hash1, hash2):
    """Tính Jaccard similarity giữa hai MinHash với độ chính xác cao."""
    if not hash1 or not hash2:
        print(f"Cảnh báo: Một trong hai hash rỗng - hash1={hash1}, hash2={hash2}")
        return 0.0  # Trả về 0 nếu một trong hai hash rỗng
    
    # Nếu hai hash giống hệt nhau, trả về 1.0 ngay lập tức
    if hash1 == hash2:
        return 1.0
    
    try:
        # Chuyển đổi từ chuỗi hash (dạng "num1,num2,num3,...") sang list số nguyên
        h1 = [int(x.strip()) for x in hash1.split(',') if x.strip()]
        h2 = [int(x.strip()) for x in hash2.split(',') if x.strip()]
        
        # Kiểm tra độ dài hash
        if len(h1) != 128 or len(h2) != 128:
            print(f"Cảnh báo: Độ dài hash không hợp lệ - len(hash1)={len(h1)}, len(hash2)={len(h2)}")
            return 0.0
        
        # Tạo MinHash từ hashvalues
        m1 = MinHash(num_perm=128)
        m2 = MinHash(num_perm=128)
        m1.hashvalues = h1
        m2.hashvalues = h2
        
        # Tính similarity với độ chính xác cao
        similarity = m1.jaccard(m2)
        
        # Debug: Kiểm tra nếu similarity thấp bất thường với hash gần giống
        if similarity < 0.7 and sum(1 for a, b in zip(h1, h2) if a == b) > 100:
            print(f"⚠️ Cảnh báo: Similarity thấp bất thường ({similarity:.4f}) dù có {sum(1 for a, b in zip(h1, h2) if a == b)} giá trị hash trùng nhau")
        
        print(f"Jaccard similarity: {similarity:.4f} (hash1={hash1[:30]}..., hash2={hash2[:30]}...)")
        return similarity
        
    except Exception as e:
        print(f"Lỗi khi tính Jaccard similarity: {str(e)}")
        return 0.0  # Trả về 0 nếu có lỗi

# def check_content_similarity(file_content, filename, current_content_hash):
#     """Kiểm tra độ giống nhau của nội dung với các tài liệu đã lưu."""
#     if not current_content_hash:
#         print(f"Bỏ qua kiểm tra độ giống cho {filename}: Không có content_hash")
#         return True, None, "Không kiểm tra độ giống nhau (không phải văn bản)"

#     blockchain.nodes = set(node_registry.get_peers())
#     print(f"Nodes trước khi đồng bộ: {blockchain.nodes}")
#     try:
#         blockchain.replace_chain()
#     except Exception as e:
#         print(f"Lỗi khi đồng bộ blockchain: {str(e)}")
#         return True, None, "Không kiểm tra độ giống (lỗi đồng bộ blockchain)"
#     print(f"Đã đồng bộ blockchain, length={len(blockchain.chain)}")

#     if len(blockchain.chain) <= 1:
#         print("Blockchain rỗng hoặc chỉ có genesis block, không kiểm tra độ giống")
#         return True, None, "Tài liệu không giống bất kỳ tài liệu nào đã lưu"

#     for block in blockchain.chain:
#         transactions = block.get('transactions', [])
#         if not transactions:
#             print(f"Block {block['index']} không có giao dịch")
#             continue
#         for tx in transactions:
#             # stored_doc_hash = tx.get('document_hash')
#             # stored_content_hash = tx.get('content_hash')
#             document_hash_obj = tx.get('document_hash', {})
#             stored_content_hash = document_hash_obj.get('content_hash')
#             stored_doc_hash = document_hash_obj.get('document_hash')
#             print(f"Kiểm tra block {block['index']}, tx: document_hash={stored_doc_hash}, content_hash={stored_content_hash}")
#             if stored_content_hash:
#                 try:
#                     similarity = jaccard_similarity(current_content_hash, stored_content_hash)
#                     if similarity >= 0.65:
#                         print(f"Tài liệu {filename} giống {similarity*100:.2f}% với hash={stored_doc_hash}")
#                         return False, stored_doc_hash, f"Tài liệu giống {similarity*100:.2f}% với hash={stored_doc_hash}"
#                 except Exception as e:
#                     print(f"Lỗi khi so sánh content_hash: {str(e)}")
#                     continue
    
#     print(f"Tài liệu {filename} không giống bất kỳ tài liệu nào đã lưu")
#     return True, None, "Tài liệu không giống bất kỳ tài liệu nào đã lưu"

def check_content_similarity(file_content, filename, current_content_hash):
    """Kiểm tra độ giống nhau của nội dung với các tài liệu đã lưu."""
    if not current_content_hash:
        print(f"Bỏ qua kiểm tra độ giống cho {filename}: Không có content_hash")
        return True, None, "Không kiểm tra độ giống nhau (không phải văn bản)"

    blockchain.nodes = set(node_registry.get_peers())
    print(f"Nodes trước khi đồng bộ: {blockchain.nodes}")
    try:
        blockchain.replace_chain()
    except Exception as e:
        print(f"Lỗi khi đồng bộ blockchain: {str(e)}")
        return True, None, "Không kiểm tra độ giống (lỗi đồng bộ blockchain)"
    print(f"Đã đồng bộ blockchain, length={len(blockchain.chain)}")

    if len(blockchain.chain) <= 1:
        print("Blockchain rỗng hoặc chỉ có genesis block, không kiểm tra độ giống")
        return True, None, "Tài liệu không giống bất kỳ tài liệu nào đã lưu"

    for block in blockchain.chain:
        transactions = block.get('transactions', [])
        if not transactions:
            print(f"Block {block['index']} không có giao dịch")
            continue
        for tx in transactions:
            document_hash_obj = tx.get('document_hash', {})
            stored_content_hash = document_hash_obj.get('content_hash')
            if stored_content_hash:
                print(f"⚠️ So sánh với content_hash lưu: {stored_content_hash}")
                try:
                    similarity = jaccard_similarity(current_content_hash, stored_content_hash)
                    if similarity >= 0.65: 
                        print(f"Tài liệu {filename} giống {similarity*100:.2f}% với hash={stored_content_hash}")
                        return False, stored_content_hash, f"Tài liệu giống {similarity*100:.2f}%"
                except Exception as e:
                    print(f"Lỗi khi so sánh content_hash: {str(e)}")
                    continue

    print(f"Tài liệu {filename} không giống bất kỳ tài liệu nào đã lưu")
    return True, None, "Tài liệu không giống bất kỳ tài liệu nào đã lưu"

@app.route('/ping', methods=['GET'])
def ping():
    return jsonify({'message': 'Node đang hoạt động'}), 200

@app.route('/get_ip', methods=['GET'])
def get_ip():
    local_ip = get_local_ip()
    port = app.config.get('PORT', 5000)
    current_node_url = f'http://{local_ip}:{port}'
    return jsonify({'ip': current_node_url}), 200

@app.route('/get_nodes', methods=['GET'])
def get_nodes():
    blockchain.nodes = set(node_registry.get_peers())
    print(f"Nodes hiện tại: {blockchain.nodes}")
    return jsonify({'nodes': list(blockchain.nodes)}), 200

@app.route('/store_document', methods=['POST'])
def store_document():
    if 'file' not in request.files:
        return jsonify({'message': 'Không có file trong request'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'message': 'Không có file được chọn'}), 400

    try:
        file_content = file.read()
        if len(file_content) == 0:
            print(f"File {file.filename} rỗng, không được phép lưu")
            return jsonify({'message': 'File rỗng', 'is_valid': False}), 400

        text = extract_text(file_content, file.filename)
        if text is None:
            print(f"File {file.filename} rỗng hoặc chỉ chứa khoảng trắng")
            return jsonify({'message': 'File rỗng hoặc chỉ chứa khoảng trắng', 'is_valid': False}), 400

        document_hash = hashlib.sha256(file_content).hexdigest()
        content_hash = get_content_hash(file_content, file.filename)

        if blockchain.verify_document(document_hash):
            print(f"Tài liệu đã tồn tại: document_hash={document_hash}")
            return jsonify({
                'message': 'Tài liệu đã tồn tại trong blockchain',
                'document_hash': document_hash,
                'is_valid': False
            }), 400

        is_text_valid, text_message = check_text_content(file_content, file.filename)
        if not is_text_valid:
            print(f"Nội dung không hợp lệ: {text_message}")
            return jsonify({'message': text_message, 'is_valid': False}), 400

        if content_hash:
            print(f"Đang kiểm tra độ giống với content_hash: {content_hash[:50]}...")
            is_content_valid, similar_hash, similarity_message = check_content_similarity(
                file_content, file.filename, content_hash
            )
            print(f"Kết quả kiểm tra: is_valid={is_content_valid}, similar_hash={similar_hash}, message={similarity_message}")
            if not is_content_valid:
                print(f"Tài liệu bị từ chối: {similarity_message}")
                return jsonify({
                    'message': similarity_message,
                    'similar_hash': similar_hash,
                    'is_valid': False
                }), 400

        data_to_send = {'document_hash': document_hash, 'content_hash': content_hash or ""}
        files_to_send = {'file': (file.filename, file_content)}
        local_response = verify_transaction_local(data_to_send, files_to_send)
        if not local_response['is_valid']:
            print(f"Kiểm tra cục bộ thất bại: {local_response['message']}")
            return jsonify(local_response), 400

        blockchain.nodes = set(node_registry.get_peers())
        current_node = f'http://{get_local_ip()}:{request.environ["SERVER_PORT"]}'
        verification_results = [local_response['is_valid']]
        print(f"Bắt đầu xác minh với nodes: {blockchain.nodes}")

        for node in blockchain.nodes:
            if node != current_node:
                response = broadcast_with_retry(
                    f'{node}/verify_transaction',
                    data=data_to_send,
                    files=files_to_send,
                    timeout=60
                )
                if response and response.status_code == 200:
                    verification_results.append(response.json()['is_valid'])
                else:
                    print(f"Node {node} không phản hồi hoặc xác minh thất bại")
                    verification_results.append(False)

        total_nodes = len(blockchain.nodes)
        valid_count = sum(1 for result in verification_results if result)
        if valid_count <= total_nodes / 2:
            print(f"Không đạt đồng thuận: {valid_count}/{total_nodes} node xác minh hợp lệ")
            return jsonify({
                'message': 'Không đạt đồng thuận',
                'valid_count': valid_count,
                'total_nodes': total_nodes
            }), 403

        transaction_data = {'document_hash': document_hash, 'content_hash': content_hash or ""}
        block_index = blockchain.add_transaction(transaction_data)
        print(f"Đã thêm giao dịch: document_hash={document_hash}, content_hash={content_hash[:50] if content_hash else 'None'}..., index={block_index}")

        for node in blockchain.nodes:
            if node != current_node:
                broadcast_with_retry(f'{node}/add_transaction', transaction_data)

        if len(blockchain.transactions) >= 1:
            previous_block = blockchain.get_previous_block()
            proof = blockchain.proof_of_work(previous_block['proof'])
            previous_hash = blockchain.hash_block(previous_block)
            new_block = blockchain.create_block(proof, previous_hash)

            for node in blockchain.nodes:
                if node != current_node:
                    broadcast_with_retry(f'{node}/add_block', new_block, timeout=60)

        print(f"Tài liệu đã lưu thành công: document_hash={document_hash}, block_index={block_index}")
        return jsonify({
            'message': 'Tài liệu đã được lưu',
            'file_hash': document_hash,
            'block_index': block_index
        }), 201

    except Exception as e:
        print(f"Lỗi khi lưu tài liệu: {str(e)}")
        return jsonify({'message': 'Lỗi khi lưu tài liệu', 'error': str(e)}), 500

def verify_transaction_local(data, files):
    """Kiểm tra giao dịch cục bộ."""
    document_hash = data.get('document_hash')
    content_hash = data.get('content_hash')
    file_content = files['file'][1]
    filename = files['file'][0]

    try:
        if len(file_content) == 0:
            print(f"File {filename} rỗng, không được phép lưu")
            return {'message': 'File rỗng', 'is_valid': False}

        text = extract_text(file_content, filename)
        if text is None:
            print(f"File {filename} rỗng hoặc chỉ chứa khoảng trắng")
            return {'message': 'File rỗng hoặc chỉ chứa khoảng trắng', 'is_valid': False}

        calculated_hash = hashlib.sha256(file_content).hexdigest()
        if calculated_hash != document_hash:
            print(f"Hash không khớp: calculated={calculated_hash}, received={document_hash}")
            return {'message': 'Hash không khớp', 'is_valid': False}

        is_text_valid, text_message = check_text_content(file_content, filename)
        if not is_text_valid:
            print(f"Nội dung không hợp lệ: {text_message}")
            return {'message': text_message, 'is_valid': False}

        if blockchain.verify_document(document_hash):
            print(f"Tài liệu đã tồn tại: document_hash={document_hash}")
            return {'message': 'Tài liệu đã tồn tại', 'is_valid': False}

        if content_hash:
            print(f"Kiểm tra độ giống cục bộ với content_hash: {content_hash[:50]}...")
            is_content_valid, similar_hash, similarity_message = check_content_similarity(
                file_content, filename, content_hash
            )
            if not is_content_valid:
                print(f"Tài liệu bị từ chối: {similarity_message}")
                return {
                    'message': similarity_message,
                    'is_valid': False,
                    'similar_hash': similar_hash
                }

        print(f"Kiểm tra cục bộ thành công: document_hash={document_hash}")
        return {'message': 'Xác minh thành công', 'is_valid': True, 'document_hash': document_hash}

    except Exception as e:
        print(f"Lỗi khi kiểm tra cục bộ: {str(e)}")
        return {'message': 'Lỗi khi xác minh', 'is_valid': False, 'error': str(e)}

@app.route('/verify_transaction', methods=['POST'])
def verify_transaction():
    if 'file' not in request.files:
        return jsonify({'message': 'Không có file trong request', 'is_valid': False}), 400
    if 'document_hash' not in request.form:
        return jsonify({'message': 'Thiếu document_hash', 'is_valid': False}), 400

    file = request.files['file']
    document_hash = request.form['document_hash']
    content_hash = request.form.get('content_hash')

    try:
        file_content = file.read()

        if len(file_content) == 0:
            print(f"File {file.filename} rỗng, không được phép lưu")
            return jsonify({'message': 'File rỗng', 'is_valid': False}), 400

        text = extract_text(file_content, file.filename)
        if text is None:
            print(f"File {file.filename} rỗng hoặc chỉ chứa khoảng trắng")
            return jsonify({'message': 'File rỗng hoặc chỉ chứa khoảng trắng', 'is_valid': False}), 400

        calculated_hash = hashlib.sha256(file_content).hexdigest()
        if calculated_hash != document_hash:
            print(f"Hash không khớp: calculated={calculated_hash}, received={document_hash}")
            return jsonify({'message': 'Hash không khớp', 'is_valid': False}), 400

        is_text_valid, text_message = check_text_content(file_content, file.filename)
        if not is_text_valid:
            print(f"Nội dung không hợp lệ: {text_message}")
            return jsonify({'message': text_message, 'is_valid': False}), 400

        if blockchain.verify_document(document_hash):
            print(f"Tài liệu đã tồn tại: document_hash={document_hash}")
            return jsonify({'message': 'Tài liệu đã tồn tại', 'is_valid': False}), 400

        if content_hash:
            print(f"Đang kiểm tra độ giống với content_hash: {content_hash[:50]}...")
            is_content_valid, similar_hash, similarity_message = check_content_similarity(
                file_content, file.filename, content_hash
            )
            print(f"Kết quả kiểm tra: is_valid={is_content_valid}, similar_hash={similar_hash}, message={similarity_message}")
            if not is_content_valid:
                print(f"Tài liệu bị từ chối: {similarity_message}")
                return jsonify({
                    'message': similarity_message,
                    'is_valid': False,
                    'similar_hash': similar_hash
                }), 400

        print(f"Tài liệu được xác minh: document_hash={document_hash}")
        return jsonify({
            'message': 'Xác minh thành công',
            'is_valid': True,
            'document_hash': document_hash
        }), 200

    except Exception as e:
        print(f"Lỗi khi xác minh tài liệu: {str(e)}")
        return jsonify({
            'message': 'Lỗi khi xác minh',
            'is_valid': False,
            'error': str(e)
        }), 500

@app.route('/store_on_ethereum', methods=['POST'])
def store_on_ethereum():
    if 'file' not in request.files:
        return jsonify({'message': 'Không có file trong request'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'message': 'Không có file được chọn'}), 400

    try:
        file_content = file.read()
        document_hash = hashlib.sha256(file_content).hexdigest()
        private_key = request.form.get('private_key')
        if not private_key:
            return jsonify({'message': 'Thiếu private key'}), 400

        account = Account.from_key(private_key)
        transaction = contract.functions.storeDocument(document_hash).build_transaction({
            'chainId': 31337,
            'gas': 100000,
            'gasPrice': web3.to_wei('10', 'gwei'),
            'nonce': web3.eth.get_transaction_count(account.address),
        })

        signed_txn = web3.eth.account.sign_transaction(transaction, private_key)
        tx_hash = web3.eth.send_raw_transaction(signed_txn.raw_transaction)
        receipt = web3.eth.wait_for_transaction_receipt(tx_hash)

        return jsonify({
            'message': 'Đã lưu trên Ethereum',
            'file_hash': document_hash,
            'tx_hash': tx_hash.hex(),
            'block_number': receipt['blockNumber']
        }), 201

    except Exception as e:
        return jsonify({'message': 'Lỗi khi lưu trên Ethereum', 'error': str(e)}), 500

@app.route('/verify_on_ethereum', methods=['POST'])
def verify_on_ethereum():
    if 'file' not in request.files:
        return jsonify({'message': 'Không có file trong request'}), 400

    file = request.files['file']
    try:
        file_content = file.read()
        document_hash = hashlib.sha256(file_content).hexdigest()
        is_stored = contract.functions.verifyDocument(document_hash).call()

        return jsonify({
            'document_hash': document_hash,
            'is_verified': is_stored,
            'message': 'Tài liệu hợp lệ' if is_stored else 'Tài liệu không tồn tại trên Ethereum'
        }), 200

    except Exception as e:
        return jsonify({'message': 'Lỗi khi xác minh', 'error': str(e)}), 500

@app.route('/add_transaction', methods=['POST'])
def add_transaction():
    json_data = request.get_json()
    if not json_data or 'document_hash' not in json_data:
        return jsonify({'message': 'Thiếu document_hash'}), 400
    index = blockchain.add_transaction(json_data)
    return jsonify({'message': f'Giao dịch sẽ được ghi vào block {index}'}), 201

@app.route('/verify_document', methods=['POST'])
def verify_document():
    if 'file' not in request.files:
        return jsonify({'message': 'Không có file trong request'}), 400

    file = request.files['file']
    try:
        file_content = file.read()
        document_hash = hashlib.sha256(file_content).hexdigest()
        is_verified = blockchain.verify_document(document_hash)

        return jsonify({
            'document_hash': document_hash,
            'is_verified': is_verified,
            'message': 'Tài liệu hợp lệ' if is_verified else 'Tài liệu không tồn tại hoặc đã bị thay đổi'
        }), 200

    except Exception as e:
        return jsonify({'message': 'Lỗi khi kiểm tra', 'error': str(e)}), 500


@app.route('/get_chain', methods=['GET'])
def get_chain():
    blockchain.nodes = set(node_registry.get_peers())
    blockchain.replace_chain()
    return jsonify({
        'chain': blockchain.chain,
        'length': len(blockchain.chain)
    }), 200
    
@app.route('/get_chain_ethereum', methods=['GET'])
def get_chain_ethereum():
    try:
        # Kết nối với contract
        contract_instance = web3.eth.contract(address=contract_address, abi=contract_abi)

        # Tạo bộ lọc cho sự kiện DocumentStored
        event_signature_hash = web3.keccak(text="DocumentStored(string)").hex()

        event_filter = {
            'fromBlock': 0,
            'toBlock': 'latest',
            'address': contract_address,
            'topics': [event_signature_hash]
        }

        # Truy vấn logs
        logs = web3.eth.get_logs(event_filter)

        # Xử lý dữ liệu logs thành chuỗi khối
        chain = []
        previous_hash = "0"  # Hash khởi đầu
        proof = 1  # Giá trị proof khởi đầu

        for index, log in enumerate(logs, start=1):
            # Giải mã dữ liệu sự kiện
            document_hash = decode(['string'], log['data'])[0]
            block_number = log['blockNumber']
            tx_hash = log['transactionHash'].hex()
            tx = web3.eth.get_transaction(tx_hash)
            sender = tx['from']
            timestamp = web3.eth.get_block(block_number)['timestamp']  # Lấy timestamp từ block

            # Tạo giao dịch từ document_hash
            transactions = [
                {
                    "document_hash": {
                        # "content_hash": "15589848778398159345",  # Giá trị giả lập, có thể thay bằng logic khác
                        "document_hash": document_hash
                    }
                }
            ]

            # Tính previous_hash (giả lập, có thể dùng hash của block trước)
            current_hash = web3.keccak(hexstr=tx_hash).hex()  # Giả lập hash hiện tại
            if index > 1:
                previous_hash = current_hash  # Hash của block trước

            # Tạo block cho chuỗi
            block = {
                "index": index,
                "previous_hash": previous_hash,
                "proof": proof,  # Giá trị proof giả lập, có thể thay bằng logic PoW
                "timestamp": timestamp,
                "transactions": transactions
            }
            chain.append(block)
            proof += 1  # Tăng proof cho block tiếp theo

        return jsonify({
            "chain": chain
        }), 200

    except Exception as e:
        return jsonify({
            'message': 'Lỗi khi truy xuất danh sách tài liệu',
            'error': str(e)
        }), 500
        
@app.route('/mine_block', methods=['POST'])
def mine_block():
    previous_block = blockchain.get_previous_block()
    proof = blockchain.proof_of_work(previous_block['proof'])
    previous_hash = blockchain.hash_block(previous_block)
    block = blockchain.create_block(proof, previous_hash)

    current_node = f'http://{get_local_ip()}:{request.environ["SERVER_PORT"]}'
    for node in blockchain.nodes:
        if node != current_node:
            broadcast_with_retry(f'{node}/add_block', block, timeout=60)

    return jsonify(block), 200

@app.route('/add_block', methods=['POST'])
def add_block():
    block = request.get_json()
    previous_block = blockchain.get_previous_block()

    if block['previous_hash'] != blockchain.hash_block(previous_block):
        return jsonify({'message': 'Block không hợp lệ: previous_hash không khớp'}), 400

    if not blockchain.is_valid_proof(block['proof'], previous_block['proof']):
        return jsonify({'message': 'Block không hợp lệ: proof không hợp lệ'}), 400

    if block['index'] != previous_block['index'] + 1:
        return jsonify({'message': 'Block không hợp lệ: index không đúng'}), 400

    blockchain.chain.append(block)
    return jsonify({'message': 'Block đã được thêm vào chain'}), 200

@app.route('/register_node', methods=['POST'])
def register_node():
    json_data = request.get_json()
    node_url = json_data.get('node_url')

    if not node_url:
        return jsonify({'message': 'Thiếu node_url'}), 400

    if not node_url.startswith('http://') and not node_url.startswith('https://'):
        node_url = f'http://{node_url}'

    if not check_node_status(node_url):
        return jsonify({'message': f'Node {node_url} không sẵn sàng'}), 503

    node_url = node_registry.register_node(node_url)
    blockchain.add_node(node_url)
    blockchain.nodes = set(node_registry.get_peers())

    current_node = f'http://{get_local_ip()}:{request.environ["SERVER_PORT"]}'
    for peer in blockchain.nodes:
        if peer != node_url and peer != current_node:
            broadcast_with_retry(f'{peer}/add_node', {'node_url': node_url})

    try:
        response = broadcast_with_retry(f'{node_url}/sync_chain', {'chain': blockchain.chain}, timeout=60)
        if response and response.status_code == 200:
            print(f"Node mới {node_url} đã đồng bộ chuỗi")
    except Exception as e:
        print(f"Lỗi khi đồng bộ với node mới: {str(e)}")

    return jsonify({
        'message': 'Node đã được đăng ký',
        'total_nodes': list(blockchain.nodes)
    }), 201

@app.route('/add_node', methods=['POST'])
def add_node():
    json_data = request.get_json()
    node_url = json_data.get('node_url')

    if not node_url:
        return jsonify({'message': 'Thiếu node_url'}), 400

    node_url = node_registry.register_node(node_url)
    blockchain.add_node(node_url)
    blockchain.nodes = set(node_registry.get_peers())

    return jsonify({'message': f'Node {node_url} đã được thêm'}), 201

@app.route('/sync_chain', methods=['POST'])
def sync_chain():
    data = request.get_json()
    if not data or 'chain' not in data:
        return jsonify({'message': 'Thiếu chuỗi trong request'}), 400

    new_chain = data['chain']
    if not blockchain.is_chain_valid(new_chain):
        return jsonify({'message': 'Chuỗi không hợp lệ'}), 400

    if len(new_chain) > len(blockchain.chain):
        blockchain.chain = new_chain
        return jsonify({'message': 'Chuỗi đã được đồng bộ'}), 200
    else:
        return jsonify({'message': 'Chuỗi không cần đồng bộ'}), 200

if __name__ == '__main__':
    port = int(sys.argv[1]) if len(sys.argv) > 1 else 5000
    app.config['PORT'] = port
    local_ip = get_local_ip()
    current_node_url = f'http://{local_ip}:{port}'

    node_registry.register_node(current_node_url)
    if bootstrap_url != current_node_url:
        try:
            response = requests.post(f'{bootstrap_url}/register_node', json={'node_url': current_node_url}, timeout=10)
            print(f"Đã đăng ký với bootstrap: {response.status_code}")
        except Exception as e:
            print(f"Lỗi khi đăng ký với bootstrap: {str(e)}")

    # node_registry.start_status_check()
    app.run(host='0.0.0.0', port=port)



