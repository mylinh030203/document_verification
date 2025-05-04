from flask import Flask, jsonify, request
import requests
import os
import hashlib
import json
import time
from blockchain import Blockchain
from web3 import Web3
from eth_account import Account
from p2p import NodeRegistry
import sys
import socket
import pdfplumber
from docx import Document
import io
from simhash import Simhash
from eth_abi import decode
from flask_cors import CORS



def get_local_ip():
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:
        s.connect(('8.8.8.8', 1))
        ip = s.getsockname()[0]
    except:
        ip = '127.0.0.1'
    finally:
        s.close()
    return ip

bootstrap_url = "http://192.168.1.8:5000"
node_registry = NodeRegistry(bootstrap_url=bootstrap_url)

# Kết nối với Ganache
web3 = Web3(Web3.HTTPProvider("http://127.0.0.1:8545"))

# Địa chỉ contract
contract_address = "0x5FbDB2315678afecb367f032d93F642f64180aa3"

# ABI của smart contract
contract_abi = [
    {
        "anonymous": False,
        "inputs": [{"indexed": False, "internalType": "string", "name": "documentHash", "type": "string"}],
        "name": "DocumentStored",
        "type": "event"
    },
    {
        "inputs": [{"internalType": "string", "name": "documentHash", "type": "string"}],
        "name": "storeDocument",
        "outputs": [],
        "stateMutability": "nonpayable",
        "type": "function"
    },
    {
        "inputs": [{"internalType": "string", "name": "documentHash", "type": "string"}],
        "name": "verifyDocument",
        "outputs": [{"internalType": "bool", "name": "", "type": "bool"}],
        "stateMutability": "view",
        "type": "function"
    }
]

# Load contract
contract = web3.eth.contract(address=contract_address, abi=contract_abi)

if web3.is_connected():
    print("✅ Đã kết nối với Ethereum node!")
else:
    print("❌ Không thể kết nối với Ethereum node!")

app = Flask(__name__)
CORS(app)
blockchain = Blockchain()

# Đảm bảo thư mục lưu trữ file
UPLOAD_FOLDER = 'Uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Danh sách từ khóa không lành mạnh
BLACKLIST_KEYWORDS = ['offensive', 'inappropriate', 'hate', 'violence', 'illegal']

def broadcast_with_retry(url, data=None, files=None, retries=3, timeout=60):
    for attempt in range(retries):
        try:
            if files:
                response = requests.post(url, data=data, files=files, timeout=timeout)
            else:
                response = requests.post(url, json=data, timeout=timeout)
            print(f"Broadcast tới {url}: {response.status_code}, response={response.text}")
            return response
        except Exception as e:
            print(f"Lỗi khi broadcast tới {url}, lần thử {attempt + 1}/{retries}: {str(e)}")
            if attempt < retries - 1:
                time.sleep(2)
    return None

def check_node_status(node_url, retries=5, delay=2):
    for attempt in range(retries):
        try:
            response = requests.get(f'{node_url}/ping', timeout=5)
            if response.status_code == 200:
                print(f"Node {node_url} sẵn sàng")
                return True
        except Exception as e:
            print(f"Node {node_url} chưa sẵn sàng, lần thử {attempt + 1}/{retries}: {str(e)}")
            time.sleep(delay)
    return False

def extract_text(file_content, filename):
    """Trích xuất văn bản từ file."""
    text_extensions = ['.txt', '.md']
    supported_extensions = text_extensions + ['.pdf', '.docx']
    
    if any(filename.lower().endswith(ext) for ext in supported_extensions):
        try:
            text = ""
            if filename.lower().endswith('.txt') or filename.lower().endswith('.md'):
                text = file_content.decode('utf-8', errors='ignore')
            elif filename.lower().endswith('.pdf'):
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
            elif filename.lower().endswith('.docx'):
                doc = Document(io.BytesIO(file_content))
                for para in doc.paragraphs:
                    text += para.text + "\n"
            return text.strip() or None
        except Exception as e:
            print(f"Lỗi khi trích xuất văn bản: {str(e)}")
            return None
    return None

def check_text_content(file_content, filename):
    """Kiểm tra nội dung văn bản để phát hiện từ khóa không lành mạnh."""
    text = extract_text(file_content, filename)
    if text:
        for keyword in BLACKLIST_KEYWORDS:
            if keyword.lower() in text.lower():
                print(f"Phát hiện từ khóa không lành mạnh: {keyword}")
                return False, f"Nội dung chứa từ khóa không lành mạnh: {keyword}"
        return True, "Nội dung văn bản hợp lệ"
    return True, "Không kiểm tra nội dung (định dạng không được hỗ trợ hoặc không trích xuất được)"

def get_content_hash(file_content, filename):
    """Tạo SimHash từ nội dung văn bản."""
    text = extract_text(file_content, filename)
    if text:
        return str(Simhash(text).value)
    return None

def hamming_distance(hash1, hash2):
    """Tính Hamming distance giữa hai hash."""
    if not hash1 or not hash2:
        return float('inf')
    return bin(int(hash1) ^ int(hash2)).count('1')

def check_content_similarity(file_content, filename, current_content_hash):
    """Kiểm tra độ giống nhau của nội dung với các tài liệu đã lưu."""
    if not current_content_hash:
        return True, None, "Không kiểm tra độ giống nhau (không phải văn bản)"

    for block in blockchain.chain:
        for tx in block.get('transactions', []):
            stored_doc_hash = tx.get('document_hash')
            stored_content_hash = tx.get('content_hash')
            if stored_content_hash:
                distance = hamming_distance(current_content_hash, stored_content_hash)
                if distance <= 12:  # Giống ≥80%
                    return False, stored_doc_hash, f"Tài liệu giống {(1 - distance/64)*100:.2f}% với hash={stored_doc_hash}"
    
    return True, None, "Tài liệu không giống bất kỳ tài liệu nào đã lưu"

@app.route('/ping', methods=['GET'])
def ping():
    return jsonify({'message': 'Node đang hoạt động'}), 200

@app.route('/get_ip', methods=['GET'])
def get_ip():
    local_ip = get_local_ip()
    port = app.config.get('PORT', 5000)
    current_node_url = f'http://{local_ip}:{port}'
    return jsonify({'ip': current_node_url}), 200

@app.route('/get_nodes', methods=['GET'])
def get_nodes():
    blockchain.nodes = set(node_registry.get_peers())
    print(f"Nodes hiện tại: {blockchain.nodes}")
    return jsonify({'nodes': list(blockchain.nodes)}), 200

# @app.route('/store_document', methods=['POST'])
# def store_document():
#     if 'file' not in request.files:
#         return jsonify({'message': 'Không có file trong request'}), 400

#     file = request.files['file']
#     if file.filename == '':
#         return jsonify({'message': 'Không có file được chọn'}), 400

#     try:
#         file_content = file.read()
#         document_hash = hashlib.sha256(file_content).hexdigest()
#         content_hash = get_content_hash(file_content, file.filename)

#         # Xác minh với các node khác
#         blockchain.nodes = set(node_registry.get_peers())
#         current_node = f'http://{get_local_ip()}:{request.environ["SERVER_PORT"]}'
#         verification_results = []
#         data_to_send = {'document_hash': document_hash}
#         if content_hash:
#             data_to_send['content_hash'] = content_hash
#         files_to_send = {'file': (file.filename, file_content)}

#         for node in blockchain.nodes:
#             if node != current_node:
#                 response = broadcast_with_retry(
#                     f'{node}/verify_transaction',
#                     data=data_to_send,
#                     files=files_to_send,
#                     timeout=60
#                 )
#                 if response and response.status_code == 200:
#                     verification_results.append(response.json()['is_valid'])

#         # Kiểm tra đồng thuận
#         total_nodes = len(blockchain.nodes) - 1
#         valid_count = sum(1 for result in verification_results if result)
#         if total_nodes > 0 and valid_count <= total_nodes / 2:
#             return jsonify({
#                 'message': 'Không đạt đồng thuận',
#                 'valid_count': valid_count,
#                 'total_nodes': total_nodes
#             }), 403

#         # Thêm giao dịch
#         transaction_data = {'document_hash': document_hash}
#         if content_hash:
#             transaction_data['content_hash'] = content_hash
#         block_index = blockchain.add_transaction(transaction_data)

#         # Broadcast giao dịch
#         for node in blockchain.nodes:
#             if node != current_node:
#                 broadcast_with_retry(f'{node}/add_transaction', transaction_data)

#         # Tạo block nếu đủ giao dịch
#         if len(blockchain.transactions) >= 1:
#             previous_block = blockchain.get_previous_block()
#             proof = blockchain.proof_of_work(previous_block['proof'])
#             previous_hash = blockchain.hash_block(previous_block)
#             new_block = blockchain.create_block(proof, previous_hash)

#             for node in blockchain.nodes:
#                 if node != current_node:
#                     broadcast_with_retry(f'{node}/add_block', new_block, timeout=60)

#         return jsonify({
#             'message': 'Tài liệu đã được lưu',
#             'file_hash': document_hash,
#             'block_index': block_index
#         }), 201

#     except Exception as e:
#         return jsonify({'message': 'Lỗi khi lưu tài liệu', 'error': str(e)}), 500

@app.route('/store_document', methods=['POST'])
def store_document():
    if 'file' not in request.files:
        return jsonify({'message': 'Không có file trong request'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'message': 'Không có file được chọn'}), 400

    try:
        file_content = file.read()
        document_hash = hashlib.sha256(file_content).hexdigest()
        content_hash = get_content_hash(file_content, file.filename)

        # Kiểm tra nếu document_hash đã tồn tại
        if blockchain.verify_document(document_hash):
            return jsonify({
                'message': 'Tài liệu đã tồn tại trong blockchain',
                'document_hash': document_hash,
                'is_valid': False
            }), 400

        # Kiểm tra độ giống nhau của nội dung
        if content_hash:
            print(f"Đang kiểm tra độ giống nhau với hash: {content_hash}")
            is_content_valid, similar_hash, similarity_message = check_content_similarity(
                file_content, file.filename, content_hash
            )
            print(f"Đã kiểm tra độ giống nhau với hash: {similar_hash}, độ giống nhau: {similarity_message}")

            if not is_content_valid:
                print(f"Tài liệu giống với hash: {similar_hash}, độ giống nhau: {similarity_message}")  
                return jsonify({
                    'message': similarity_message,
                    'similar_hash': similar_hash,
                    'is_valid': False
                }), 400

        # Xác minh với các node khác
        blockchain.nodes = set(node_registry.get_peers())
        current_node = f'http://{get_local_ip()}:{request.environ["SERVER_PORT"]}'
        verification_results = []
        data_to_send = {'document_hash': document_hash}
        if content_hash:
            data_to_send['content_hash'] = content_hash
        files_to_send = {'file': (file.filename, file_content)}

        for node in blockchain.nodes:
            if node != current_node:
                response = broadcast_with_retry(
                    f'{node}/verify_transaction',
                    data=data_to_send,
                    files=files_to_send,
                    timeout=60
                )
                if response and response.status_code == 200:
                    verification_results.append(response.json()['is_valid'])

        # Kiểm tra đồng thuận
        total_nodes = len(blockchain.nodes) - 1
        valid_count = sum(1 for result in verification_results if result)
        if total_nodes > 0 and valid_count <= total_nodes / 2:
            return jsonify({
                'message': 'Không đạt đồng thuận',
                'valid_count': valid_count,
                'total_nodes': total_nodes
            }), 403

        # Thêm giao dịch
        transaction_data = {'document_hash': document_hash}
        if content_hash:
            transaction_data['content_hash'] = content_hash
        block_index = blockchain.add_transaction(transaction_data)

        # Broadcast giao dịch
        for node in blockchain.nodes:
            if node != current_node:
                broadcast_with_retry(f'{node}/add_transaction', transaction_data)

        # Tạo block nếu đủ giao dịch
        if len(blockchain.transactions) >= 1:
            previous_block = blockchain.get_previous_block()
            proof = blockchain.proof_of_work(previous_block['proof'])
            previous_hash = blockchain.hash_block(previous_block)
            new_block = blockchain.create_block(proof, previous_hash)

            for node in blockchain.nodes:
                if node != current_node:
                    broadcast_with_retry(f'{node}/add_block', new_block, timeout=60)

        return jsonify({
            'message': 'Tài liệu đã được lưu',
            'file_hash': document_hash,
            'block_index': block_index
        }), 201

    except Exception as e:
        return jsonify({'message': 'Lỗi khi lưu tài liệu', 'error': str(e)}), 500
@app.route('/verify_transaction', methods=['POST'])
def verify_transaction():
    if 'file' not in request.files:
        return jsonify({'message': 'Không có file trong request', 'is_valid': False}), 400
    if 'document_hash' not in request.form:
        return jsonify({'message': 'Thiếu document_hash', 'is_valid': False}), 400

    file = request.files['file']
    document_hash = request.form['document_hash']
    content_hash = request.form.get('content_hash')

    try:
        file_content = file.read()

        # Kiểm tra 1: File rỗng
        if len(file_content) == 0:
            return jsonify({'message': 'File rỗng', 'is_valid': False}), 400

        # Kiểm tra 2: Mã hash
        calculated_hash = hashlib.sha256(file_content).hexdigest()
        if calculated_hash != document_hash:
            return jsonify({'message': 'Hash không khớp', 'is_valid': False}), 400

        # Kiểm tra 3: Nội dung văn bản
        is_text_valid, text_message = check_text_content(file_content, file.filename)
        if not is_text_valid:
            return jsonify({'message': text_message, 'is_valid': False}), 400

        # Kiểm tra 4: Trùng lặp
        if blockchain.verify_document(document_hash):
            return jsonify({'message': 'Tài liệu đã tồn tại', 'is_valid': False}), 400

        # Kiểm tra 5: Độ giống nhau
        is_content_valid, similar_hash, similarity_message = check_content_similarity(
            file_content, file.filename, content_hash)
        if not is_content_valid:
            return jsonify({
                'message': similarity_message,
                'is_valid': False,
                'similar_hash': similar_hash
            }), 400

        return jsonify({
            'message': 'Xác minh thành công',
            'is_valid': True,
            'document_hash': document_hash
        }), 200

    except Exception as e:
        return jsonify({
            'message': 'Lỗi khi xác minh',
            'is_valid': False,
            'error': str(e)
        }), 500

@app.route('/store_on_ethereum', methods=['POST'])
def store_on_ethereum():
    if 'file' not in request.files:
        return jsonify({'message': 'Không có file trong request'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'message': 'Không có file được chọn'}), 400

    try:
        file_content = file.read()
        document_hash = hashlib.sha256(file_content).hexdigest()
        private_key = request.form.get('private_key')
        if not private_key:
            return jsonify({'message': 'Thiếu private key'}), 400

        account = Account.from_key(private_key)
        transaction = contract.functions.storeDocument(document_hash).build_transaction({
            'chainId': 31337,
            'gas': 100000,
            'gasPrice': web3.to_wei('10', 'gwei'),
            'nonce': web3.eth.get_transaction_count(account.address),
        })

        signed_txn = web3.eth.account.sign_transaction(transaction, private_key)
        tx_hash = web3.eth.send_raw_transaction(signed_txn.raw_transaction)
        receipt = web3.eth.wait_for_transaction_receipt(tx_hash)

        return jsonify({
            'message': 'Đã lưu trên Ethereum',
            'file_hash': document_hash,
            'tx_hash': tx_hash.hex(),
            'block_number': receipt['blockNumber']
        }), 201

    except Exception as e:
        return jsonify({'message': 'Lỗi khi lưu trên Ethereum', 'error': str(e)}), 500

@app.route('/verify_on_ethereum', methods=['POST'])
def verify_on_ethereum():
    if 'file' not in request.files:
        return jsonify({'message': 'Không có file trong request'}), 400

    file = request.files['file']
    try:
        file_content = file.read()
        document_hash = hashlib.sha256(file_content).hexdigest()
        is_stored = contract.functions.verifyDocument(document_hash).call()

        return jsonify({
            'document_hash': document_hash,
            'is_verified': is_stored,
            'message': 'Tài liệu hợp lệ' if is_stored else 'Tài liệu không tồn tại trên Ethereum'
        }), 200

    except Exception as e:
        return jsonify({'message': 'Lỗi khi xác minh', 'error': str(e)}), 500

@app.route('/add_transaction', methods=['POST'])
def add_transaction():
    json_data = request.get_json()
    if not json_data or 'document_hash' not in json_data:
        return jsonify({'message': 'Thiếu document_hash'}), 400
    index = blockchain.add_transaction(json_data)
    return jsonify({'message': f'Giao dịch sẽ được ghi vào block {index}'}), 201

@app.route('/verify_document', methods=['POST'])
def verify_document():
    if 'file' not in request.files:
        return jsonify({'message': 'Không có file trong request'}), 400

    file = request.files['file']
    try:
        file_content = file.read()
        document_hash = hashlib.sha256(file_content).hexdigest()
        is_verified = blockchain.verify_document(document_hash)

        return jsonify({
            'document_hash': document_hash,
            'is_verified': is_verified,
            'message': 'Tài liệu hợp lệ' if is_verified else 'Tài liệu không tồn tại hoặc đã bị thay đổi'
        }), 200

    except Exception as e:
        return jsonify({'message': 'Lỗi khi kiểm tra', 'error': str(e)}), 500

@app.route('/get_chain', methods=['GET'])
def get_chain():
    blockchain.nodes = set(node_registry.get_peers())
    blockchain.replace_chain()
    return jsonify({
        'chain': blockchain.chain,
        'length': len(blockchain.chain)
    }), 200

@app.route('/get_chain_ethereum', methods=['GET'])
def get_chain_ethereum():
    try:
        # Kết nối với contract
        contract_instance = web3.eth.contract(address=contract_address, abi=contract_abi)

        # Tạo bộ lọc cho sự kiện DocumentStored
        event_signature_hash = web3.keccak(text="DocumentStored(string)").hex()

        event_filter = {
            'fromBlock': 0,
            'toBlock': 'latest',
            'address': contract_address,
            'topics': [event_signature_hash]
        }

        # Truy vấn logs
        logs = web3.eth.get_logs(event_filter)

        # Xử lý dữ liệu logs thành chuỗi khối
        chain = []
        previous_hash = "0"  # Hash khởi đầu
        proof = 1  # Giá trị proof khởi đầu

        for index, log in enumerate(logs, start=1):
            # Giải mã dữ liệu sự kiện
            document_hash = decode(['string'], log['data'])[0]
            block_number = log['blockNumber']
            tx_hash = log['transactionHash'].hex()
            tx = web3.eth.get_transaction(tx_hash)
            sender = tx['from']
            timestamp = web3.eth.get_block(block_number)['timestamp']  # Lấy timestamp từ block

            # Tạo giao dịch từ document_hash
            transactions = [
                {
                    "document_hash": {
                        # "content_hash": "15589848778398159345",  # Giá trị giả lập, có thể thay bằng logic khác
                        "document_hash": document_hash
                    }
                }
            ]

            # Tính previous_hash (giả lập, có thể dùng hash của block trước)
            current_hash = web3.keccak(hexstr=tx_hash).hex()  # Giả lập hash hiện tại
            if index > 1:
                previous_hash = current_hash  # Hash của block trước

            # Tạo block cho chuỗi
            block = {
                "index": index,
                "previous_hash": previous_hash,
                "proof": proof,  # Giá trị proof giả lập, có thể thay bằng logic PoW
                "timestamp": timestamp,
                "transactions": transactions
            }
            chain.append(block)
            proof += 1  # Tăng proof cho block tiếp theo

        return jsonify({
            "chain": chain
        }), 200

    except Exception as e:
        return jsonify({
            'message': 'Lỗi khi truy xuất danh sách tài liệu',
            'error': str(e)
        }), 500
# @app.route('/get_chain_ethereum', methods=['GET'])
# def get_chain_ethereum():
#     try:
#         # Kết nối với contract
#         contract_instance = web3.eth.contract(address=contract_address, abi=contract_abi)

#         # Tạo bộ lọc cho sự kiện DocumentStored
#         event_signature_hash = web3.keccak(text="DocumentStored(string)").hex()

#         event_filter = {
#             'fromBlock': 0,
#             'toBlock': 'latest',
#             'address': contract_address,
#             'topics': [event_signature_hash]
#         }

#         # Truy vấn logs
#         logs = web3.eth.get_logs(event_filter)

#         # Xử lý dữ liệu logs
#         documents = []
#         for log in logs:
#             # Giải mã dữ liệu sự kiện
#             document_hash = decode(['string'], log['data'])[0] 
#             block_number = log['blockNumber']
#             tx_hash = log['transactionHash'].hex()
#             # Lấy sender từ transaction
#             tx = web3.eth.get_transaction(tx_hash)
#             sender = tx['from']

#             documents.append({
#                 'document_hash': document_hash,
#                 'sender': sender,
#                 'block_number': block_number,
#                 'tx_hash': tx_hash
#             })

#         return jsonify({
#             'message': 'Danh sách tài liệu đã lưu trên Ethereum',
#             'documents': documents,
#             'total': len(documents)
#         }), 200

#     except Exception as e:
#         return jsonify({
#             'message': 'Lỗi khi truy xuất danh sách tài liệu',
#             'error': str(e)
#         }), 500
        
@app.route('/mine_block', methods=['POST'])
def mine_block():
    previous_block = blockchain.get_previous_block()
    proof = blockchain.proof_of_work(previous_block['proof'])
    previous_hash = blockchain.hash_block(previous_block)
    block = blockchain.create_block(proof, previous_hash)

    current_node = f'http://{get_local_ip()}:{request.environ["SERVER_PORT"]}'
    for node in blockchain.nodes:
        if node != current_node:
            broadcast_with_retry(f'{node}/add_block', block, timeout=60)

    return jsonify(block), 200

@app.route('/add_block', methods=['POST'])
def add_block():
    block = request.get_json()
    previous_block = blockchain.get_previous_block()

    if block['previous_hash'] != blockchain.hash_block(previous_block):
        return jsonify({'message': 'Block không hợp lệ: previous_hash không khớp'}), 400

    if not blockchain.is_valid_proof(block['proof'], previous_block['proof']):
        return jsonify({'message': 'Block không hợp lệ: proof không hợp lệ'}), 400

    if block['index'] != previous_block['index'] + 1:
        return jsonify({'message': 'Block không hợp lệ: index không đúng'}), 400

    blockchain.chain.append(block)
    return jsonify({'message': 'Block đã được thêm vào chain'}), 200

@app.route('/register_node', methods=['POST'])
def register_node():
    json_data = request.get_json()
    node_url = json_data.get('node_url')

    if not node_url:
        return jsonify({'message': 'Thiếu node_url'}), 400

    if not node_url.startswith('http://') and not node_url.startswith('https://'):
        node_url = f'http://{node_url}'

    if not check_node_status(node_url):
        return jsonify({'message': f'Node {node_url} không sẵn sàng'}), 503

    node_url = node_registry.register_node(node_url)
    blockchain.add_node(node_url)
    blockchain.nodes = set(node_registry.get_peers())

    current_node = f'http://{get_local_ip()}:{request.environ["SERVER_PORT"]}'
    for peer in blockchain.nodes:
        if peer != node_url and peer != current_node:
            broadcast_with_retry(f'{peer}/add_node', {'node_url': node_url})

    try:
        response = broadcast_with_retry(f'{node_url}/sync_chain', {'chain': blockchain.chain}, timeout=60)
        if response and response.status_code == 200:
            print(f"Node mới {node_url} đã đồng bộ chuỗi")
    except Exception as e:
        print(f"Lỗi khi đồng bộ với node mới: {str(e)}")

    return jsonify({
        'message': 'Node đã được đăng ký',
        'total_nodes': list(blockchain.nodes)
    }), 201

@app.route('/add_node', methods=['POST'])
def add_node():
    json_data = request.get_json()
    node_url = json_data.get('node_url')

    if not node_url:
        return jsonify({'message': 'Thiếu node_url'}), 400

    node_url = node_registry.register_node(node_url)
    blockchain.add_node(node_url)
    blockchain.nodes = set(node_registry.get_peers())

    return jsonify({'message': f'Node {node_url} đã được thêm'}), 201

@app.route('/sync_chain', methods=['POST'])
def sync_chain():
    data = request.get_json()
    if not data or 'chain' not in data:
        return jsonify({'message': 'Thiếu chuỗi trong request'}), 400

    new_chain = data['chain']
    if not blockchain.is_chain_valid(new_chain):
        return jsonify({'message': 'Chuỗi không hợp lệ'}), 400

    if len(new_chain) > len(blockchain.chain):
        blockchain.chain = new_chain
        return jsonify({'message': 'Chuỗi đã được đồng bộ'}), 200
    else:
        return jsonify({'message': 'Chuỗi không cần đồng bộ'}), 200

if __name__ == '__main__':
    port = int(sys.argv[1]) if len(sys.argv) > 1 else 5000
    app.config['PORT'] = port
    local_ip = get_local_ip()
    current_node_url = f'http://{local_ip}:{port}'

    node_registry.register_node(current_node_url)
    if bootstrap_url != current_node_url:
        try:
            response = requests.post(f'{bootstrap_url}/register_node', json={'node_url': current_node_url}, timeout=10)
            print(f"Đã đăng ký với bootstrap: {response.status_code}")
        except Exception as e:
            print(f"Lỗi khi đăng ký với bootstrap: {str(e)}")

    # node_registry.start_status_check()
    app.run(host='0.0.0.0', port=port)